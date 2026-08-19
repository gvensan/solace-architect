#!/usr/bin/env python3
"""
Generate a Solace Architect intake DOCX with Structured Document Tags (SDTs).

Usage:
    python3 scripts/build-intake-docx.py --output intake.docx
    python3 scripts/build-intake-docx.py --output filled.docx --data intake.yaml

Modes:
    (no --data)   Produces a blank intake template with placeholder text and
                  dropdown prompts.  Give this to the customer to fill out.

    --data <yaml> Reads a YAML file in the same format produced by
                  parse-intake-docx.py and populates every SDT from it.
                  The output round-trips through the parser without loss.

The SDT alias/tag pairs mirror FIELD_MAP in parse-intake-docx.py exactly.
"""

import argparse
import os
import sys
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is required.  Install with: pip install python-docx")
    sys.exit(1)

try:
    from lxml import etree
except ImportError:
    print("ERROR: lxml is required.  Install with: pip install lxml")
    sys.exit(1)

try:
    import yaml
    HAS_YAML = True
except ImportError:
    HAS_YAML = False

# ---------------------------------------------------------------------------
# Namespace constant (same as parser)
# ---------------------------------------------------------------------------
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ---------------------------------------------------------------------------
# FIELD_MAP — mirrors parse-intake-docx.py FIELD_MAP exactly
# ---------------------------------------------------------------------------
FIELD_MAP = {
    # S1: Contact
    'contact_name':  'contact.name',
    'contact_role':  'contact.role',
    'contact_email': 'contact.email',
    'contact_phone': 'contact.phone',
    'contact_org':   'contact.organization',
    'contact_date':  'contact.date',

    # S2: Project
    'project_name':   'project.name',
    'project_type':   'project.type',
    's2_other':       'project.notes',
    's2_references':  'project.references',

    # S3: System Landscape
    'existing_messaging': 'landscape.existing_messaging',
    'protocols_in_use':   'landscape.protocols_in_use',
    'aggregate_volumes':  'landscape.volumes',
    'schemas':            'landscape.schemas',
    'industry_vertical':  'landscape.vertical',
    'vertical_other':     'landscape.vertical_other',
    's3_other':           'landscape.notes',
    's3_references':      'landscape.references',

    # S4: Domain — Banking
    'banking_regulatory': 'domain.banking.regulatory_constraints',
    'banking_messaging':  'domain.banking.existing_messaging_infrastructure',
    'banking_auth':       'domain.banking.authorization_model',
    'banking_data_class': 'domain.banking.data_classification',
    'banking_audience':   'domain.banking.internal_vs_customer_facing',

    # S4: Domain — Capital Markets
    'capmkts_latency':    'domain.capital_markets.latency_budget',
    'capmkts_topology':   'domain.capital_markets.global_topology',
    'capmkts_feeds':      'domain.capital_markets.feed_infrastructure',
    'capmkts_messaging':  'domain.capital_markets.existing_messaging',
    'capmkts_compliance': 'domain.capital_markets.compliance_and_replay',

    # S4: Domain — Manufacturing
    'mfg_ot_protocols': 'domain.manufacturing.ot_protocol_inventory',
    'mfg_edge':         'domain.manufacturing.edge_constraints',
    'mfg_telemetry':    'domain.manufacturing.telemetry_vs_command',
    'mfg_historians':   'domain.manufacturing.existing_historians_mes',

    # S4: Domain — Healthcare
    'hc_hipaa':   'domain.healthcare.hipaa_phi',
    'hc_interop': 'domain.healthcare.interoperability_standards',
    'hc_realtime':'domain.healthcare.realtime_vs_batch',

    # S4: Domain — Retail
    'retail_systems':    'domain.retail.order_inventory_systems',
    'retail_peak':       'domain.retail.peak_traffic',
    'retail_omnichannel':'domain.retail.omnichannel',
    'retail_realtime':   'domain.retail.personalization_realtime',

    # S4: Domain — Telecom
    'telco_events':  'domain.telecom.network_event_types',
    'telco_oss_bss': 'domain.telecom.oss_bss_integration',
    'telco_scale':   'domain.telecom.subscriber_scale',
    'telco_5g':      'domain.telecom.five_g_edge',

    # S4: Domain — Logistics
    'logistics_tracking': 'domain.logistics.tracking_visibility',
    'logistics_partners': 'domain.logistics.partner_integration',
    'logistics_warehouse':'domain.logistics.warehouse_systems',
    'logistics_custody':  'domain.logistics.chain_of_custody',

    # S4: Domain — Energy
    'energy_scada': 'domain.energy.scada_grid',
    'energy_meters':'domain.energy.smart_meters',
    'energy_nerc':  'domain.energy.regulatory_nerc',
    'energy_der':   'domain.energy.renewable_der',

    # S4: Domain — Government
    'gov_classification': 'domain.government.classification_level',
    'gov_interagency':    'domain.government.interagency_sharing',
    'gov_citizen':        'domain.government.citizen_services',
    'gov_compliance':     'domain.government.compliance_frameworks',

    # S4: Domain - Other. A catch-all vertical still has designable
    # specifics; `domain.notes` below stays as the free-form addendum.
    'other_regulatory': 'domain.other.regulatory_constraints',
    'other_sensitivity':'domain.other.data_sensitivity',
    'other_platform':   'domain.other.platform_constraints',
    'other_rules':      'domain.other.domain_rules',

    's4_other':      'domain.notes',
    's4_references': 'domain.references',

    # S5: Migration
    'source_platform':       'migration.source_platform',
    'source_platform_other': 'migration.source_platform_other',
    'source_version':        'migration.source_version',
    'migration_artifact_count': 'migration.artifact_count',
    'migration_app_count':   'migration.app_count',
    'migration_driver':      'migration.driver',
    'coexistence':           'migration.coexistence_strategy',
    'coexistence_duration':  'migration.coexistence_duration',
    'migration_constraints': 'migration.constraints',
    'data_migration':        'migration.data_migration',
    's5_other':              'migration.notes',
    's5_references':         'migration.references',

    # S6: SAM
    'sam_use_case':     'sam.use_case',
    'llm_provider':     'sam.llm_provider',
    'sam_agent_count':  'sam.agent_count',
    'sam_capabilities': 'sam.capabilities',
    'sam_backends':     'sam.backends',
    'sam_channels':     'sam.channels',
    'sam_channels_list':'sam.channels_list',
    'sam_existing_ai':  'sam.existing_ai',
    'sam_volume':       'sam.volume',
    's6_other':         'sam.notes',
    's6_references':    'sam.references',

    # S7: Security
    'auth_method':       'security.authentication',
    'tls':               'security.tls',
    'encryption_at_rest':'security.encryption_at_rest',
    'security_network':  'security.network_isolation',
    'security_acl':      'security.access_control',
    'security_audit':    'security.audit_compliance',
    'security_gateway':  'security.api_gateway',
    'security_secrets':  'security.secret_management',
    's7_other':          'security.notes',
    's7_references':     'security.references',

    # S8: Technical Requirements
    'delivery_mode':        'requirements.delivery_mode',
    'ordering':             'requirements.ordering',
    'processing_guarantee': 'requirements.processing_guarantee',
    'latency_tier':         'requirements.latency_tier',
    'topology':             'requirements.topology',
    'sites_regions':        'requirements.sites_and_regions',
    'it_ot_boundary':       'requirements.it_ot_boundary',
    'growth':               'requirements.growth_expectations',
    'data_residency':       'requirements.data_residency',
    'ops_team':             'requirements.operations_team',
    'solace_experience':    'requirements.solace_experience',
    'observability':        'requirements.observability',
    'cicd':                 'requirements.cicd',
    's8_other':             'requirements.notes',
    's8_references':        'requirements.references',

    # S9: Goals
    'goal_driver':           'goals.driver',
    'goal_timeline':         'goals.timeline',
    'goal_budget':           'goals.budget',
    'goal_team_size':        'goals.team_size',
    'goal_org_constraints':  'goals.organizational_constraints',
    's9_other':              'goals.notes',
    's9_references':         'goals.references',

    # S10: Preferences
    'execution_mode':          'preferences.execution_mode',
    'provision_event_portal':  'preferences.provision_event_portal',

    # S11: Additional
    'additional_notes':      'additional.notes',
    'additional_references': 'additional.references',
}


# ---------------------------------------------------------------------------
# Dropdown option definitions
# ---------------------------------------------------------------------------
PROJECT_TYPE_OPTIONS = [
    ("Select one...", ""),
    ("New Build",     "new_build"),
    ("Migration",     "migration"),
    ("Extension",     "extension"),
    ("SAM Integration", "sam"),
]

INDUSTRY_VERTICAL_OPTIONS = [
    ("Select one...", ""),
    ("Banking",           "banking"),
    ("Capital Markets",   "capital_markets"),
    ("Manufacturing",     "manufacturing"),
    ("Healthcare",        "healthcare"),
    ("Retail",            "retail"),
    ("Telecom",           "telecom"),
    ("Logistics",         "logistics"),
    ("Energy",            "energy"),
    ("Government",        "government"),
    ("Other",             "other"),
]

SOURCE_PLATFORM_OPTIONS = [
    ("Select one...", ""),
    ("IBM MQ",    "ibm_mq"),
    ("Kafka",     "kafka"),
    ("RabbitMQ",  "rabbitmq"),
    ("TIBCO",     "tibco"),
    ("Other",     "other"),
]

COEXISTENCE_OPTIONS = [
    ("Select one...", ""),
    ("Big bang", "big_bang"),
    ("Phased",   "phased"),
    ("Bridge",   "bridge"),
]

LLM_PROVIDER_OPTIONS = [
    ("Select one...",  ""),
    ("OpenAI",         "openai"),
    ("Anthropic",      "anthropic"),
    ("Google",         "google"),
    ("AWS Bedrock",    "bedrock"),
    ("Azure OpenAI",   "azure"),
    ("Self-hosted",    "self_hosted"),
    ("Other",          "other"),
]

SAM_CHANNELS_OPTIONS = [
    ("Select one...", ""),
    ("Web chat",  "web"),
    ("Slack",     "slack"),
    ("Teams",     "teams"),
    ("Mobile",    "mobile"),
    ("API",       "api"),
    ("Multiple",  "multiple"),
]

AUTH_METHOD_OPTIONS = [
    ("Select one...",      ""),
    ("OAuth 2.0 / OIDC",  "oauth"),
    ("SAML",              "saml"),
    ("mTLS",              "mtls"),
    ("Basic auth",        "basic"),
    ("API keys",          "api_keys"),
    ("Multiple",          "multiple"),
]

TLS_OPTIONS = [
    ("Select one...",        ""),
    ("Required everywhere",  "required"),
    ("External only",        "external_only"),
    ("Not required",         "not_required"),
]

ENCRYPTION_AT_REST_OPTIONS = [
    ("Select one...",    ""),
    ("Required",         "required"),
    ("Not required",     "not_required"),
    ("Under evaluation", "evaluating"),
]

DELIVERY_MODE_OPTIONS = [
    ("Select one...", ""),
    ("Direct",      "direct"),
    ("Guaranteed",  "guaranteed"),
    ("Mixed",       "mixed"),
]

ORDERING_OPTIONS = [
    ("Select one...", ""),
    ("None",     "none"),
    ("Per-key",  "per_key"),
    ("Global",   "global"),
]

PROCESSING_GUARANTEE_OPTIONS = [
    ("Select one...",    ""),
    ("At-least-once",  "at_least_once"),
    ("At-most-once",   "at_most_once"),
]

LATENCY_TIER_OPTIONS = [
    ("Select one...",           ""),
    ("Sub-millisecond (<1ms)",  "sub_millisecond"),
    ("Sub-second (<1s)",        "sub_second"),
    ("Seconds (1-10s)",         "seconds"),
    ("Minutes",                 "minutes"),
]

TOPOLOGY_OPTIONS = [
    ("Select one...", ""),
    ("Single site",   "single_site"),
    ("Multi-region",  "multi_region"),
    ("Hybrid cloud",  "hybrid_cloud"),
    ("Edge",          "edge"),
]

EXECUTION_MODE_OPTIONS = [
    ("Select one...", ""),
    ("Auto",         "auto"),
    ("Interactive",  "interactive"),
]

PROVISION_EVENT_PORTAL_OPTIONS = [
    ("No — design-only (default)", "false"),
    ("Yes — provision into Solace Cloud after design", "true"),
]

SYSTEM_ROLE_OPTIONS = [
    ("Select one...",        ""),
    ("Producer",             "producer"),
    ("Consumer",             "consumer"),
    ("Producer / Consumer",  "producer_consumer"),
]

EVENT_DELIVERY_OPTIONS = [
    ("Select one...", ""),
    ("Direct",      "direct"),
    ("Guaranteed",  "guaranteed"),
]


# ---------------------------------------------------------------------------
# Data lookup helpers
# ---------------------------------------------------------------------------

def _get_nested(d: dict, dotted_key: str):
    """Return the value at dot-notation path in nested dict, or None."""
    parts = dotted_key.split('.')
    cur = d
    for p in parts:
        if isinstance(cur, dict) and p in cur:
            cur = cur[p]
        else:
            return None
    return cur


def _lookup(data: Optional[dict], tag: str) -> Optional[str]:
    """Look up a scalar field value by SDT tag name."""
    if data is None:
        return None
    dotted = FIELD_MAP.get(tag)
    if dotted is None:
        return None
    val = _get_nested(data, dotted)
    return str(val) if val is not None else None


def _display_for_value(options: list, raw_value: str) -> Optional[str]:
    """Given a machine value, return its display text from an options list."""
    for display, value in options:
        if value == raw_value:
            return display
    return None


def _lookup_dropdown(data: Optional[dict], tag: str, options: list) -> Optional[str]:
    """
    Look up a dropdown field.  Returns the display text matching the stored
    machine value, or None if not set.
    """
    raw = _lookup(data, tag)
    if raw is None:
        return None
    display = _display_for_value(options, raw)
    return display  # None if the value is unknown (caller treats as unset)


# ---------------------------------------------------------------------------
# XML / SDT construction helpers
# ---------------------------------------------------------------------------

def _make_el(tag: str, attribs: dict = None) -> etree._Element:
    """Create an lxml element in the W namespace."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def _make_run_xml(text: str, bold: bool = False, italic: bool = False,
                  color: Optional[str] = None, font: str = 'Calibri',
                  size_pt: int = 11) -> etree._Element:
    """Build a <w:r> element with the given text and formatting."""
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    fnt = OxmlElement('w:rFonts')
    fnt.set(qn('w:ascii'),  font)
    fnt.set(qn('w:hAnsi'), font)
    rpr.append(fnt)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size_pt * 2))
    rpr.append(sz)

    if bold:
        rpr.append(OxmlElement('w:b'))
    if italic:
        rpr.append(OxmlElement('w:i'))
    if color:
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), color)
        rpr.append(clr)

    r.append(rpr)

    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)

    return r


def _make_sdtPr(tag: str, alias: str) -> etree._Element:
    """Build <w:sdtPr> with alias and tag child elements."""
    sdtPr = OxmlElement('w:sdtPr')

    alias_el = OxmlElement('w:alias')
    alias_el.set(qn('w:val'), alias)
    sdtPr.append(alias_el)

    tag_el = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), tag)
    sdtPr.append(tag_el)

    return sdtPr


def add_text_sdt(paragraph, tag: str, alias: str,
                 placeholder: str = "Click here to enter text...",
                 value: Optional[str] = None) -> etree._Element:
    """
    Append a text content-control SDT to *paragraph*.

    When value is None the SDT shows placeholder text (showingPlcHdr set).
    When value is provided the SDT content is the real text.

    Returns the <w:sdt> element appended to paragraph._p.
    """
    sdt = OxmlElement('w:sdt')

    # sdtPr
    sdtPr = _make_sdtPr(tag, alias)
    txt_el = OxmlElement('w:text')
    if value is None:
        sdtPr.append(OxmlElement('w:showingPlcHdr'))
    sdtPr.append(txt_el)
    sdt.append(sdtPr)

    # sdtContent
    sdtContent = OxmlElement('w:sdtContent')
    inner_p = OxmlElement('w:p')

    if value is None:
        # Placeholder: gray italic
        run = _make_run_xml(placeholder, italic=True, color='808080')
    else:
        run = _make_run_xml(value)

    inner_p.append(run)
    sdtContent.append(inner_p)
    sdt.append(sdtContent)

    paragraph._p.append(sdt)
    return sdt


def add_dropdown_sdt(paragraph, tag: str, alias: str,
                     options: list,
                     value: Optional[str] = None) -> etree._Element:
    """
    Append a dropdown (combo box) content-control SDT to *paragraph*.

    options: list of (displayText, value) tuples.
    value:   the display text to pre-select; None leaves "Select one..." shown.

    Returns the <w:sdt> element appended to paragraph._p.
    """
    sdt = OxmlElement('w:sdt')

    # sdtPr
    sdtPr = _make_sdtPr(tag, alias)

    ddl = OxmlElement('w:dropDownList')
    for display_text, val in options:
        li = OxmlElement('w:listItem')
        li.set(qn('w:displayText'), display_text)
        li.set(qn('w:value'), val)
        ddl.append(li)
    sdtPr.append(ddl)
    sdt.append(sdtPr)

    # sdtContent
    sdtContent = OxmlElement('w:sdtContent')
    inner_p = OxmlElement('w:p')

    display = value if value is not None else options[0][0]
    color = None if value is not None else '808080'
    run = _make_run_xml(display, color=color)

    inner_p.append(run)
    sdtContent.append(inner_p)
    sdt.append(sdtContent)

    paragraph._p.append(sdt)
    return sdt


def add_cell_text_sdt(cell, row_idx: int, col_key: str,
                      placeholder: str = "Enter value...",
                      value: Optional[str] = None) -> etree._Element:
    """
    Insert a text SDT into a table cell.  The tag is derived from
    col_key + row index (e.g. system_name_r1) so the parser's skip logic
    is triggered correctly for repeated-row cells.
    """
    tag = f"{col_key}_r{row_idx}"
    alias = col_key.replace('_', ' ').title()
    paragraph = cell.paragraphs[0]
    return add_text_sdt(paragraph, tag, alias, placeholder, value)


def add_cell_dropdown_sdt(cell, row_idx: int, col_key: str,
                          options: list,
                          value: Optional[str] = None) -> etree._Element:
    """Insert a dropdown SDT into a table cell."""
    tag = f"{col_key}_r{row_idx}"
    alias = col_key.replace('_', ' ').title()
    paragraph = cell.paragraphs[0]
    return add_dropdown_sdt(paragraph, tag, alias, options, value)


# ---------------------------------------------------------------------------
# Document-level style helpers
# ---------------------------------------------------------------------------

def _set_default_font(doc: Document, font_name: str = 'Calibri', size_pt: int = 11):
    """Set the Normal style default font and size."""
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(size_pt)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a heading paragraph."""
    doc.add_heading(text, level=level)


def _add_label(doc: Document, text: str) -> None:
    """Add a bold label paragraph with light-gray background shading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

    # Paragraph shading (light gray)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)

    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)


def _add_field_text(doc: Document, tag: str, alias: str,
                    placeholder: str, data: Optional[dict]) -> None:
    """Label + text SDT pair."""
    _add_label(doc, alias)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    value = _lookup(data, tag)
    add_text_sdt(p, tag, alias, placeholder, value)


def _add_field_dropdown(doc: Document, tag: str, alias: str,
                        options: list, data: Optional[dict]) -> None:
    """Label + dropdown SDT pair."""
    _add_label(doc, alias)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    display = _lookup_dropdown(data, tag, options)
    add_dropdown_sdt(p, tag, alias, options, display)


# ---------------------------------------------------------------------------
# Table builders
# ---------------------------------------------------------------------------

def _style_header_cell(cell, text: str) -> None:
    """Bold, dark-blue-shaded header cell."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F497D')
    tcPr.append(shd)


def _add_cell_borders(table):
    """Add thin borders to every cell in the table."""
    tbl = table._tbl
    for tr in tbl.findall(f'{{{W}}}tr'):
        for tc in tr.findall(f'{{{W}}}tc'):
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'),   'single')
                border.set(qn('w:sz'),    '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'BFBFBF')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def _build_systems_table(doc: Document, data: Optional[dict]) -> None:
    """
    Add the Systems table (S3).

    Columns: System Name | Role | Protocol | Owner
    Parser expects headers exactly as below (case-insensitive 'system name').
    """
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    headers = ['System Name', 'Role', 'Protocol', 'Owner']
    for i, h in enumerate(headers):
        _style_header_cell(table.rows[0].cells[i], h)
    _add_cell_borders(table)

    # Collect rows from data if available
    systems = []
    if data:
        raw = _get_nested(data, 'landscape.systems')
        if isinstance(raw, list):
            systems = raw

    # Add populated rows, then at least 3 blank template rows
    populated = len(systems)
    total_rows = max(populated, 3)

    for row_idx in range(1, total_rows + 1):
        row = table.add_row()
        sys_data = systems[row_idx - 1] if row_idx - 1 < populated else {}

        # Col 0: System Name (text SDT)
        add_cell_text_sdt(row.cells[0], row_idx, 'system_name',
                          'e.g. Order Management System',
                          sys_data.get('System Name'))

        # Col 1: Role (dropdown SDT)
        role_display = None
        raw_role = sys_data.get('Role')
        if raw_role:
            role_display = _display_for_value(SYSTEM_ROLE_OPTIONS, raw_role) or raw_role
        add_cell_dropdown_sdt(row.cells[1], row_idx, 'system_role',
                              SYSTEM_ROLE_OPTIONS, role_display)

        # Col 2: Protocol (text SDT)
        add_cell_text_sdt(row.cells[2], row_idx, 'system_protocol',
                          'e.g. AMQP, MQTT, REST',
                          sys_data.get('Protocol'))

        # Col 3: Owner (text SDT)
        add_cell_text_sdt(row.cells[3], row_idx, 'system_owner',
                          'e.g. Platform Team',
                          sys_data.get('Owner'))

    _add_cell_borders(table)
    doc.add_paragraph()  # breathing space after table


def _build_events_table(doc: Document, data: Optional[dict]) -> None:
    """
    Add the Events table (S3).

    Columns: Event Name | Rate | Delivery | Payload Format | Payload Size
    Parser expects header 'event name' (case-insensitive).
    Payload size feeds broker spool sizing and protocol-overhead decisions.
    """
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    headers = ['Event Name', 'Rate', 'Delivery', 'Payload Format', 'Payload Size']
    for i, h in enumerate(headers):
        _style_header_cell(table.rows[0].cells[i], h)
    _add_cell_borders(table)

    events = []
    if data:
        raw = _get_nested(data, 'landscape.events')
        if isinstance(raw, list):
            events = raw

    populated = len(events)
    total_rows = max(populated, 3)

    for row_idx in range(1, total_rows + 1):
        row = table.add_row()
        ev_data = events[row_idx - 1] if row_idx - 1 < populated else {}

        # Col 0: Event Name (text SDT)
        add_cell_text_sdt(row.cells[0], row_idx, 'event_name',
                          'e.g. OrderPlaced',
                          ev_data.get('Event Name'))

        # Col 1: Rate (text SDT)
        add_cell_text_sdt(row.cells[1], row_idx, 'event_rate',
                          'e.g. 500/s peak',
                          ev_data.get('Rate'))

        # Col 2: Delivery (dropdown SDT)
        delivery_display = None
        raw_delivery = ev_data.get('Delivery')
        if raw_delivery:
            delivery_display = _display_for_value(EVENT_DELIVERY_OPTIONS, raw_delivery) or raw_delivery
        add_cell_dropdown_sdt(row.cells[2], row_idx, 'event_delivery',
                              EVENT_DELIVERY_OPTIONS, delivery_display)

        # Col 3: Payload Format (text SDT)
        add_cell_text_sdt(row.cells[3], row_idx, 'event_payload',
                          'e.g. JSON, Avro, Protobuf',
                          ev_data.get('Payload Format'))

        # Col 4: Payload Size (text SDT) — feeds broker sizing
        add_cell_text_sdt(row.cells[4], row_idx, 'event_payload_size',
                          'e.g. ~5KB or ~50KB peak / ~2KB typical',
                          ev_data.get('Payload Size'))

    _add_cell_borders(table)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _section_divider(doc: Document) -> None:
    """Add a thin horizontal rule paragraph between sections."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4472C4')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def build_s1_contact(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S1 — Contact Information", level=1)
    _add_field_text(doc, 'contact_name',  "Full Name",       "e.g. Jane Smith",              data)
    _add_field_text(doc, 'contact_role',  "Role / Title",    "e.g. Chief Architect",          data)
    _add_field_text(doc, 'contact_email', "Email",           "e.g. jane@example.com",         data)
    _add_field_text(doc, 'contact_phone', "Phone",           "e.g. +1-555-000-0000",          data)
    _add_field_text(doc, 'contact_org',   "Organisation",    "e.g. Acme Corp",                data)
    _add_field_text(doc, 'contact_date',  "Date Completed",  "e.g. 2026-05-04",               data)
    _section_divider(doc)


def build_s2_project(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S2 — Project Overview", level=1)
    _add_field_text(    doc, 'project_name', "Project Name", "e.g. Order Event Mesh Phase 1", data)
    _add_field_dropdown(doc, 'project_type', "Project Type", PROJECT_TYPE_OPTIONS,            data)
    _add_field_text(    doc, 's2_other',     "Other Notes",  "Any additional project context", data)
    _add_field_text(    doc, 's2_references',"References",   "Links, docs, or ticket numbers", data)
    _section_divider(doc)


def build_s3_landscape(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S3 — System Landscape", level=1)
    _add_field_text(    doc, 'existing_messaging', "Existing Messaging",
                        "Describe current messaging infrastructure",               data)
    _add_field_text(    doc, 'protocols_in_use',   "Protocols in Use",
                        "e.g. REST, AMQP 1.0, MQTT 3.1",                          data)
    _add_field_text(    doc, 'aggregate_volumes',  "Aggregate Volumes",
                        "e.g. 10,000 msg/s peak, 2 TB/day",                       data)
    _add_field_text(    doc, 'schemas',            "Schema Formats",
                        "e.g. Avro with Confluent registry, JSON Schema",          data)
    _add_field_dropdown(doc, 'industry_vertical',  "Industry Vertical",
                        INDUSTRY_VERTICAL_OPTIONS,                                 data)
    _add_field_text(    doc, 'vertical_other',     "Vertical (if Other)",
                        "Describe your industry if Other was selected",            data)
    _add_field_text(    doc, 's3_other',           "Other Landscape Notes",
                        "Any other context about the system landscape",            data)
    _add_field_text(    doc, 's3_references',      "References",
                        "Architecture diagrams, runbooks, Confluence pages",       data)

    doc.add_heading("Systems", level=2)
    _build_systems_table(doc, data)

    doc.add_heading("Events / Data Flows", level=2)
    _build_events_table(doc, data)
    _section_divider(doc)


def build_s4_domain(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S4 — Domain-Specific Questions", level=1)
    p = doc.add_paragraph(
        "Complete only the subsection(s) matching your industry vertical(s). "
        "Leave other subsections blank."
    )
    p.paragraph_format.space_after = Pt(8)

    # --- Banking ---
    _add_heading(doc, "Banking", level=2)
    _add_field_text(doc, 'banking_regulatory', "Regulatory Constraints",
                    "e.g. FINRA, MiFID II, Basel III reporting obligations",   data)
    _add_field_text(doc, 'banking_messaging',  "Existing Messaging Infrastructure",
                    "Describe internal MOM landscape",                         data)
    _add_field_text(doc, 'banking_auth',       "Authorization Model",
                    "e.g. entitlements by product desk, client segmentation",  data)
    _add_field_text(doc, 'banking_data_class', "Data Classification",
                    "e.g. public, internal, confidential, restricted",         data)
    _add_field_text(doc, 'banking_audience',   "Internal vs Customer-Facing",
                    "Which flows are external-facing?",                        data)

    # --- Capital Markets ---
    _add_heading(doc, "Capital Markets", level=2)
    _add_field_text(doc, 'capmkts_latency',    "Latency Budget",
                    "e.g. <100 µs for tick data, <1 ms for order routing",    data)
    _add_field_text(doc, 'capmkts_topology',   "Global Topology",
                    "Primary and DR sites, co-lo requirements",                data)
    _add_field_text(doc, 'capmkts_feeds',      "Feed Infrastructure",
                    "e.g. Reuters, Bloomberg, direct exchange feeds",          data)
    _add_field_text(doc, 'capmkts_messaging',  "Existing Messaging",
                    "e.g. 29West, Solace appliance, IBM MQ",                  data)
    _add_field_text(doc, 'capmkts_compliance', "Compliance and Replay",
                    "Audit replay requirements, retention windows",            data)

    # --- Manufacturing ---
    _add_heading(doc, "Manufacturing", level=2)
    _add_field_text(doc, 'mfg_ot_protocols', "OT Protocol Inventory",
                    "e.g. OPC-UA, Modbus, PROFINET in use",                   data)
    _add_field_text(doc, 'mfg_edge',         "Edge Constraints",
                    "Bandwidth, connectivity, local buffering requirements",   data)
    _add_field_text(doc, 'mfg_telemetry',    "Telemetry vs Command",
                    "Ratio of sensor telemetry to control commands",           data)
    _add_field_text(doc, 'mfg_historians',   "Existing Historians / MES",
                    "e.g. OSIsoft PI, Wonderware, SAP MES",                   data)

    # --- Healthcare ---
    _add_heading(doc, "Healthcare", level=2)
    _add_field_text(doc, 'hc_hipaa',   "HIPAA / PHI Handling",
                    "Which flows carry PHI? De-identification approach?",      data)
    _add_field_text(doc, 'hc_interop', "Interoperability Standards",
                    "e.g. HL7 FHIR, HL7 v2, DICOM, X12",                     data)
    _add_field_text(doc, 'hc_realtime',"Real-time vs Batch",
                    "Which use cases need sub-second latency?",                data)

    # --- Retail ---
    _add_heading(doc, "Retail", level=2)
    _add_field_text(doc, 'retail_systems',    "Order & Inventory Systems",
                    "e.g. SAP Commerce, Oracle OMS, custom",                  data)
    _add_field_text(doc, 'retail_peak',       "Peak Traffic",
                    "e.g. Black Friday volumes, flash sale patterns",          data)
    _add_field_text(doc, 'retail_omnichannel',"Omnichannel",
                    "Online, in-store, mobile, marketplace channels",          data)
    _add_field_text(doc, 'retail_realtime',   "Personalisation / Real-time",
                    "Real-time offer engines, inventory visibility needs",     data)

    # --- Telecom ---
    _add_heading(doc, "Telecom", level=2)
    _add_field_text(doc, 'telco_events',  "Network Event Types",
                    "e.g. alarms, KPI streams, CDRs, location events",        data)
    _add_field_text(doc, 'telco_oss_bss', "OSS/BSS Integration",
                    "Which OSS/BSS systems need event integration?",           data)
    _add_field_text(doc, 'telco_scale',   "Subscriber Scale",
                    "e.g. 50 M subscribers, 200 M IoT devices",               data)
    _add_field_text(doc, 'telco_5g',      "5G / Edge",
                    "MEC, network slicing, NWDAF event feeds",                 data)

    # --- Logistics ---
    _add_heading(doc, "Logistics", level=2)
    _add_field_text(doc, 'logistics_tracking', "Tracking & Visibility",
                    "End-to-end shipment tracking architecture",               data)
    _add_field_text(doc, 'logistics_partners', "Partner Integration",
                    "Carriers, 3PLs, customs brokers",                        data)
    _add_field_text(doc, 'logistics_warehouse',"Warehouse Systems",
                    "e.g. Manhattan, Blue Yonder, SAP EWM",                   data)
    _add_field_text(doc, 'logistics_custody', "Chain of Custody",
                    "Proof-of-delivery, handoff events, disputes",            data)

    # --- Energy ---
    _add_heading(doc, "Energy", level=2)
    _add_field_text(doc, 'energy_scada',  "SCADA / Grid Systems",
                    "SCADA vendors, grid topology, control-centre setup",      data)
    _add_field_text(doc, 'energy_meters', "Smart Meters / AMI",
                    "Meter volumes, head-end system, data cadence",            data)
    _add_field_text(doc, 'energy_nerc',   "Regulatory (NERC CIP / IEC 62351)",
                    "Applicable standards and compliance obligations",         data)
    _add_field_text(doc, 'energy_der',    "Renewable / DER",
                    "Solar, wind, battery storage event patterns",             data)

    # --- Government ---
    _add_heading(doc, "Government", level=2)
    _add_field_text(doc, 'gov_classification', "Data Classification Level",
                    "e.g. Unclassified, CUI, Secret, FedRAMP boundary",       data)
    _add_field_text(doc, 'gov_interagency',    "Interagency Sharing",
                    "Which agencies share events? Cross-domain guards?",       data)
    _add_field_text(doc, 'gov_citizen',        "Citizen Services",
                    "Public-facing event flows, accessibility requirements",   data)
    _add_field_text(doc, 'gov_compliance',     "Compliance Frameworks",
                    "e.g. FedRAMP, FISMA, NIST 800-53, CJIS",                 data)

    # --- Other ---
    _add_heading(doc, "Other", level=2)
    _add_field_text(doc, 'other_regulatory',  "Regulatory Constraints",
                    "Which regimes apply, and to which data? \"None\" is a real answer",  data)
    _add_field_text(doc, 'other_sensitivity', "Data Sensitivity",
                    "What is inside the payloads, and who may see it",          data)
    _add_field_text(doc, 'other_platform',    "Platform Constraints",
                    "Operating systems, hosting, or tooling the design must live with", data)
    _add_field_text(doc, 'other_rules',       "Domain Rules A Designer Would Not Guess",
                    "Conventions or invariants the event and topic design must respect", data)

    # --- Shared ---
    _add_field_text(doc, 's4_other',      "Other Domain Notes",
                    "Anything not captured above",                             data)
    _add_field_text(doc, 's4_references', "References",
                    "Domain-specific documents or standards",                  data)
    _section_divider(doc)


def build_s5_migration(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S5 — Migration (complete if applicable)", level=1)
    _add_field_dropdown(doc, 'source_platform',       "Source Platform",
                        SOURCE_PLATFORM_OPTIONS,                              data)
    _add_field_text(    doc, 'source_platform_other', "Source Platform (if Other)",
                        "Describe the source platform",                       data)
    _add_field_text(    doc, 'source_version',         "Source Version",
                        "e.g. Kafka 3.6, IBM MQ 9.3",                        data)
    _add_field_text(    doc, 'migration_artifact_count',"Artifact Count",
                        "Queues, topics, channels to migrate",                data)
    _add_field_text(    doc, 'migration_app_count',    "Application Count",
                        "Number of producer/consumer apps",                   data)
    _add_field_text(    doc, 'migration_driver',       "Migration Driver",
                        "Why migrate now? Cost, EOL, capability gap?",        data)
    _add_field_dropdown(doc, 'coexistence',            "Coexistence Strategy",
                        COEXISTENCE_OPTIONS,                                  data)
    _add_field_text(    doc, 'coexistence_duration',   "Coexistence Duration",
                        "e.g. 6 months, until Q4 2026",                      data)
    _add_field_text(    doc, 'migration_constraints',  "Constraints",
                        "Downtime windows, frozen periods, regulatory holds",  data)
    _add_field_text(    doc, 'data_migration',         "Data Migration",
                        "Replaying historical messages? Offset mapping?",     data)
    _add_field_text(    doc, 's5_other',               "Other Migration Notes",
                        "Anything not captured above",                        data)
    _add_field_text(    doc, 's5_references',          "References",
                        "Migration guides, runbooks",                         data)
    _section_divider(doc)


def build_s6_sam(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S6 — Solace Agent Mesh / AI Integration (complete if applicable)", level=1)
    _add_field_text(    doc, 'sam_use_case',     "Use Case",
                        "Describe the AI agent or assistant use case",        data)
    _add_field_dropdown(doc, 'llm_provider',     "LLM Provider",
                        LLM_PROVIDER_OPTIONS,                                 data)
    _add_field_text(    doc, 'sam_agent_count',  "Agent Count",
                        "Estimated number of agents",                         data)
    _add_field_text(    doc, 'sam_capabilities', "Agent Capabilities",
                        "e.g. RAG, tool use, memory, streaming",              data)
    _add_field_text(    doc, 'sam_backends',     "Backend Systems",
                        "Systems agents will call: CRM, ERP, databases",      data)
    _add_field_dropdown(doc, 'sam_channels',     "Primary Channel",
                        SAM_CHANNELS_OPTIONS,                                 data)
    _add_field_text(    doc, 'sam_channels_list',"All Channels",
                        "List all channels if Multiple selected above",       data)
    _add_field_text(    doc, 'sam_existing_ai',  "Existing AI Infrastructure",
                        "Any existing ML platforms, vector stores, pipelines", data)
    _add_field_text(    doc, 'sam_volume',       "Interaction Volume",
                        "e.g. 10,000 sessions/day, 500 concurrent",           data)
    _add_field_text(    doc, 's6_other',         "Other SAM Notes",
                        "Anything not captured above",                        data)
    _add_field_text(    doc, 's6_references',    "References",
                        "SAM docs, agent mesh architecture links",            data)
    _section_divider(doc)


def build_s7_security(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S7 — Security Requirements", level=1)
    _add_field_dropdown(doc, 'auth_method',         "Authentication Method",
                        AUTH_METHOD_OPTIONS,                                  data)
    _add_field_dropdown(doc, 'tls',                 "TLS Policy",
                        TLS_OPTIONS,                                          data)
    _add_field_dropdown(doc, 'encryption_at_rest',  "Encryption at Rest",
                        ENCRYPTION_AT_REST_OPTIONS,                           data)
    _add_field_text(    doc, 'security_network',    "Network Isolation",
                        "VPC, private link, DMZ, firewall topology",          data)
    _add_field_text(    doc, 'security_acl',        "Access Control",
                        "Topic ACL model, subscription policies",             data)
    _add_field_text(    doc, 'security_audit',      "Audit & Compliance",
                        "Log retention, SIEM integration, compliance scope",  data)
    _add_field_text(    doc, 'security_gateway',    "API Gateway / Perimeter",
                        "Reverse proxy, WAF, edge security components",       data)
    _add_field_text(    doc, 'security_secrets',    "Secret Management",
                        "e.g. HashiCorp Vault, AWS Secrets Manager",         data)
    _add_field_text(    doc, 's7_other',            "Other Security Notes",
                        "Anything not captured above",                        data)
    _add_field_text(    doc, 's7_references',       "References",
                        "Security architecture docs, threat models",          data)
    _section_divider(doc)


def build_s8_requirements(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S8 — Technical Requirements", level=1)
    _add_field_dropdown(doc, 'delivery_mode',         "Delivery Mode",
                        DELIVERY_MODE_OPTIONS,                                data)
    _add_field_dropdown(doc, 'ordering',              "Message Ordering",
                        ORDERING_OPTIONS,                                     data)
    _add_field_dropdown(doc, 'processing_guarantee',  "Processing Guarantee",
                        PROCESSING_GUARANTEE_OPTIONS,                         data)
    _add_field_dropdown(doc, 'latency_tier',          "Latency Tier",
                        LATENCY_TIER_OPTIONS,                                 data)
    _add_field_dropdown(doc, 'topology',              "Deployment Topology",
                        TOPOLOGY_OPTIONS,                                     data)
    _add_field_text(    doc, 'sites_regions',         "Sites / Regions",
                        "e.g. us-east-1, eu-west-1, on-prem Chicago DC",      data)
    _add_field_text(    doc, 'it_ot_boundary',        "IT/OT Boundary",
                        "Describe if OT systems connect to this mesh",        data)
    _add_field_text(    doc, 'growth',                "Growth Expectations",
                        "e.g. 3× volume in 18 months, 2 new regions in 2027", data)
    _add_field_text(    doc, 'data_residency',        "Data Residency",
                        "Sovereignty requirements, cross-border restrictions", data)
    _add_field_text(    doc, 'ops_team',              "Operations Team",
                        "Team size, skills, on-call model",                   data)
    _add_field_text(    doc, 'solace_experience',     "Solace Experience",
                        "Prior Solace deployments, certifications, training",  data)
    _add_field_text(    doc, 'observability',         "Observability Stack",
                        "e.g. Datadog, Prometheus/Grafana, Splunk",           data)
    _add_field_text(    doc, 'cicd',                  "CI/CD Pipeline",
                        "e.g. GitHub Actions, Jenkins, ArgoCD",               data)
    _add_field_text(    doc, 's8_other',              "Other Requirement Notes",
                        "Anything not captured above",                        data)
    _add_field_text(    doc, 's8_references',         "References",
                        "NFR documents, SLAs, SLOs",                          data)
    _section_divider(doc)


def build_s9_goals(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S9 — Goals & Constraints", level=1)
    _add_field_text(doc, 'goal_driver',           "Primary Driver",
                    "What is the top business outcome this project must achieve?", data)
    _add_field_text(doc, 'goal_timeline',         "Timeline",
                    "e.g. MVP in Q3 2026, full production Q1 2027",               data)
    _add_field_text(doc, 'goal_budget',           "Budget",
                    "Approximate budget range or constraints",                      data)
    _add_field_text(doc, 'goal_team_size',        "Team Size",
                    "Number of engineers available for this project",              data)
    _add_field_text(doc, 'goal_org_constraints',  "Organisational Constraints",
                    "Procurement policies, approved vendor lists, change windows", data)
    _add_field_text(doc, 's9_other',              "Other Goal Notes",
                    "Anything not captured above",                                 data)
    _add_field_text(doc, 's9_references',         "References",
                    "Business case docs, OKRs",                                    data)
    _section_divider(doc)


def build_s10_preferences(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S10 — Engagement Preferences", level=1)
    _add_field_dropdown(doc, 'execution_mode', "Execution Mode",
                        EXECUTION_MODE_OPTIONS,                    data)
    p = doc.add_paragraph(
        "Auto: Solace Architect runs all applicable skills and presents findings at the end.  "
        "Interactive: pauses after each skill for your input before proceeding."
    )
    p.paragraph_format.space_after = Pt(6)

    _add_field_dropdown(doc, 'provision_event_portal',
                        "Provision Event Portal model into Solace Cloud after design?",
                        PROVISION_EVENT_PORTAL_OPTIONS,            data)
    p = doc.add_paragraph(
        "When set to Yes, /solace-ep-provision runs after Event Portal design and materializes the catalog "
        "(application domains, schemas, events, applications) directly into your Solace Cloud tenant via the "
        "Solace Event Portal Designer MCP. Requires the EP Designer MCP installed in your AI host and a "
        "Solace Cloud API token with Event Portal Designer Read+Write permission. If the MCP is not "
        "configured at run time, the step records a BLOCKED status with the exact reason — it never writes "
        "silently or skips silently. Leave as No for design-only engagements that do not touch your tenant."
    )
    p.paragraph_format.space_after = Pt(6)
    _section_divider(doc)


def build_s11_additional(doc: Document, data: Optional[dict]) -> None:
    _add_heading(doc, "S11 — Additional Information", level=1)
    _add_field_text(doc, 'additional_notes',      "Additional Notes",
                    "Anything not captured in the sections above",                 data)
    _add_field_text(doc, 'additional_references', "Additional References",
                    "Any other relevant links, documents, or contact details",     data)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------

def build_title_page(doc: Document) -> None:
    """Add a simple title block at the start of the document."""
    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run("Solace Architect — Intake Form")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark blue

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    run2 = sub_p.add_run("Pre-Engagement Information Gathering")
    run2.font.name = 'Calibri'
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    # Instructions block
    doc.add_paragraph()
    instr = doc.add_paragraph()
    instr.paragraph_format.space_after = Pt(12)
    instr_run = instr.add_run(
        "Instructions: Complete the fields below before your Solace Architect engagement.  "
        "Click each blue field to enter text, or use the dropdowns to select an option.  "
        "Sections S4–S6 are conditional — fill only those relevant to your project.  "
        "When complete, save this file and share it with your Solace Architect contact, "
        "or run: python3 scripts/parse-intake-docx.py filled-intake.docx to import it directly."
    )
    instr_run.font.name = 'Calibri'
    instr_run.font.size = Pt(10)
    instr_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Decorative rule
    rule_p = doc.add_paragraph()
    pPr = rule_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F497D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    rule_p.paragraph_format.space_after = Pt(16)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_intake_docx(output_path: str, data: Optional[dict] = None) -> None:
    """
    Build the complete intake DOCX and write it to output_path.

    data: parsed YAML dict (same structure as parse-intake-docx.py output).
          Pass None to produce a blank template.
    """
    doc = Document()
    _set_default_font(doc)

    # Remove the default empty paragraph Word inserts
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    build_title_page(doc)
    build_s1_contact(doc, data)
    build_s2_project(doc, data)
    build_s3_landscape(doc, data)
    build_s4_domain(doc, data)
    build_s5_migration(doc, data)
    build_s6_sam(doc, data)
    build_s7_security(doc, data)
    build_s8_requirements(doc, data)
    build_s9_goals(doc, data)
    build_s10_preferences(doc, data)
    build_s11_additional(doc, data)

    doc.save(output_path)
    print(f"Intake DOCX written to: {output_path}")
    if data is not None:
        print("(Populated from data file)")
    else:
        print("(Blank template — share with customer to fill out)")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def _load_data(yaml_path: str) -> dict:
    """Load and validate the YAML data file."""
    if not os.path.exists(yaml_path):
        print(f"ERROR: Data file not found: {yaml_path}")
        sys.exit(1)

    with open(yaml_path, 'r', encoding='utf-8') as f:
        content = f.read()

    if HAS_YAML:
        data = yaml.safe_load(content)
    else:
        import json
        data = json.loads(content)

    if not isinstance(data, dict):
        print("ERROR: Data file must be a YAML/JSON object (dict) at the top level.")
        sys.exit(1)

    # Strip internal _meta key produced by the parser — not needed for generation
    data.pop('_meta', None)
    return data


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a Solace Architect intake DOCX with SDT content controls.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        '--output', '-o',
        required=True,
        metavar='PATH',
        help="Path to write the output DOCX (required).",
    )
    parser.add_argument(
        '--data', '-d',
        metavar='YAML_FILE',
        default=None,
        help="YAML file produced by parse-intake-docx.py to pre-populate the form.",
    )
    args = parser.parse_args()

    data = None
    if args.data:
        data = _load_data(args.data)

    output_dir = os.path.dirname(os.path.abspath(args.output))
    if not os.path.isdir(output_dir):
        print(f"ERROR: Output directory does not exist: {output_dir}")
        sys.exit(1)

    build_intake_docx(args.output, data)


if __name__ == '__main__':
    main()
